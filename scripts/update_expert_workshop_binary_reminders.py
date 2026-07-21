"""Create the v6.2 expert-workshop DOCX with printable tables and binary reminder outcomes."""

from __future__ import annotations

import io
import os
import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source materials" / "VPrivCal_Expert_Workshop_Interface_Aligned_v6_1.docx"
OUTPUT = ROOT / "source materials" / "VPrivCal_Expert_Workshop_Interface_Aligned_v6_2.docx"
QA_ASSETS = ROOT / ".docx_qa" / "v6_2_assets"
TABLE_WIDTH_TWIPS = 10080
TABLE_INDENT_TWIPS = 144


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    filename = "arialbd.ttf" if bold else "arial.ttf"
    return ImageFont.truetype(str(Path(os.environ["WINDIR"]) / "Fonts" / filename), size)


def wrapped(draw: ImageDraw.ImageDraw, text: str, max_width: int, text_font: ImageFont.FreeTypeFont) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=text_font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_multiline(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    max_width: int,
    text_font: ImageFont.FreeTypeFont,
    fill: str,
    spacing: int = 8,
) -> int:
    x, y = xy
    lines = wrapped(draw, text, max_width, text_font)
    line_height = text_font.size + spacing
    for line in lines:
        draw.text((x, y), line, font=text_font, fill=fill)
        y += line_height
    return y


def workflow_diagram(path: Path) -> None:
    image = Image.new("RGB", (2400, 760), "white")
    draw = ImageDraw.Draw(image)
    navy, gray = "#173F5F", "#66758A"
    draw.text((80, 42), "VPrivCal workflow: binary privacy-reminder calibration", font=font(48, True), fill=navy)
    draw.text(
        (80, 104),
        "Q10 and Probe calibrate a two-action rule; held-out clips test immediate acceptance.",
        font=font(25),
        fill=gray,
    )
    boxes = [
        ("1", "VPrivCal-Q10", "Set six category defaults and four reminder-trigger preferences.", "~2 min", "#E8F0F7", navy),
        ("2", "VPrivCal-Probe", "Review three private, public, or semi-public scenes and choose reminder / no reminder.", "~3–5 min", "#F0EAFB", "#6E43B9"),
        ("3", "Binary policy", "Combine Q10 priors with category-by-scene corrections into SHOW or NO REMINDER.", "Immediate", "#E6F3EF", "#27836C"),
        ("4", "Held-out videos", "Apply generic, Q10-only, and full VPrivCal policies to unseen cues.", "Evaluation", "#FFF0E7", "#C65312"),
        ("5", "Acceptance", "Collect 1–5 acceptance, preferred binary decision, false reminders, and missed reminders.", "Primary outcomes", "#E7F3F6", "#18788A"),
    ]
    left, top, width, height, gap = 70, 195, 405, 385, 65
    for index, (number, title, body, footer, bg, accent) in enumerate(boxes):
        x = left + index * (width + gap)
        draw.rounded_rectangle((x, top, x + width, top + height), radius=24, fill=bg, outline=accent, width=5)
        draw.ellipse((x + 24, top + 25, x + 86, top + 87), fill=accent)
        draw.text((x + 46, top + 39), number, font=font(27, True), fill="white", anchor="mm")
        draw.text((x + 105, top + 38), title, font=font(28, True), fill=accent)
        draw_multiline(draw, (x + 26, top + 128), body, width - 52, font(24), "#26384A", 9)
        draw.text((x + 26, top + height - 55), footer, font=font(24, True), fill=accent)
        if index < len(boxes) - 1:
            arrow_x = x + width + 12
            center_y = top + height // 2
            draw.polygon(
                [(arrow_x, center_y - 19), (arrow_x + 36, center_y), (arrow_x, center_y + 19)],
                fill=navy,
            )
            draw.rectangle((arrow_x - 12, center_y - 7, arrow_x + 4, center_y + 7), fill=navy)
    draw.rounded_rectangle((85, 625, 2315, 710), radius=18, fill="#F3F6F8", outline="#C6D1DB", width=3)
    draw.text(
        (1200, 667),
        "Interpret results as immediate cue-level acceptance—not long-term behavior or real-world effectiveness.",
        font=font(25, True),
        fill=navy,
        anchor="mm",
    )
    image.save(path, quality=96)


def comparison_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 650), "white")
    draw = ImageDraw.Draw(image)
    navy, gray = "#173F5F", "#66758A"
    draw.text((65, 35), "Held-out comparison: binary reminder decisions", font=font(43, True), fill=navy)
    draw.text(
        (65, 91),
        "Counterbalance conditions and use the same manually verified cue windows.",
        font=font(24),
        fill=gray,
    )
    boxes = [
        ("Generic policy", "One fixed reminder rule for all participants.", "#E8F0F7", "#173F5F"),
        ("Q10-only policy", "Binary category priors plus cross-cutting trigger preferences.", "#F0EAFB", "#6E43B9"),
        ("Full VPrivCal policy", "Q10 priors corrected by scene-specific Probe responses.", "#E6F3EF", "#27836C"),
    ]
    width, height, top, gap = 500, 265, 175, 85
    for i, (title, body, bg, accent) in enumerate(boxes):
        x = 65 + i * (width + gap)
        draw.rounded_rectangle((x, top, x + width, top + height), radius=24, fill=bg, outline=accent, width=5)
        draw.text((x + 28, top + 34), title, font=font(30, True), fill=accent)
        draw_multiline(draw, (x + 28, top + 105), body, width - 56, font(25), "#26384A", 10)
    draw.rounded_rectangle((65, 490, 1735, 603), radius=20, fill="#FFF5EA", outline="#D37A26", width=4)
    draw.text(
        (900, 530),
        "Each cue produces SHOW REMINDER or NO REMINDER.",
        font=font(27, True),
        fill="#8B4712",
        anchor="mm",
    )
    draw.text(
        (900, 568),
        "Participant rates acceptance (1–5) and states the preferred binary decision.",
        font=font(24),
        fill="#8B4712",
        anchor="mm",
    )
    image.save(path, quality=96)


def replace_text_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text
    for old, new in replacements.items():
        if old in text:
            text = text.replace(old, new)
    if text != paragraph.text:
        paragraph.text = text


def set_table_geometry(table) -> None:
    table.autofit = False
    table.alignment = None
    properties = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:jc"):
        for node in properties.findall(qn(tag)):
            properties.remove(node)
    width = OxmlElement("w:tblW")
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(TABLE_WIDTH_TWIPS))
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_TWIPS))
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    justification = OxmlElement("w:jc")
    justification.set(qn("w:val"), "left")
    properties.append(width)
    properties.append(indent)
    properties.append(layout)
    properties.append(justification)

    grid = table._tbl.tblGrid
    grid_columns = list(grid)
    original = [int(column.get(qn("w:w"), "1")) for column in grid_columns]
    original_total = sum(original) or len(original)
    scaled = [max(120, round(TABLE_WIDTH_TWIPS * value / original_total)) for value in original]
    scaled[-1] += TABLE_WIDTH_TWIPS - sum(scaled)
    for column, value in zip(grid_columns, scaled):
        column.set(qn("w:w"), str(value))

    for row_index, row in enumerate(table.rows):
        tr_properties = row._tr.get_or_add_trPr()
        for height in tr_properties.findall(qn("w:trHeight")):
            tr_properties.remove(height)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_properties.append(repeat)
        for column_index, cell in enumerate(row.cells):
            cell.width = Inches(scaled[min(column_index, len(scaled) - 1)] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(scaled[min(column_index, len(scaled) - 1)]))
            margins = cell_properties.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side in ("top", "left", "bottom", "right"):
                side_node = margins.find(qn(f"w:{side}"))
                if side_node is None:
                    side_node = OxmlElement(f"w:{side}")
                    margins.append(side_node)
                side_node.set(qn("w:w"), "72")
                side_node.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    if len(table.columns) >= 5:
                        run.font.size = Pt(8)


def main() -> None:
    QA_ASSETS.mkdir(parents=True, exist_ok=True)
    workflow_path = QA_ASSETS / "workflow_binary.png"
    comparison_path = QA_ASSETS / "comparison_binary.png"
    workflow_diagram(workflow_path)
    comparison_diagram(comparison_path)

    document = Document(SOURCE)
    replacements = {
        "Are the revised questions, implemented three-scene Probe interface and Hint walkthrough, workflow, and held-out evaluation sufficiently clear, actionable, low-burden, and safe to proceed to user cognitive interviews?":
            "Are the revised binary reminder questions, three-scene Probe interface, Hint walkthrough, and immediate cue-level acceptance evaluation sufficiently clear, actionable, low-burden, and safe for user cognitive interviews?",
        "Version 6.1 | July 2026 | Interface-aligned edition":
            "Version 6.2 | July 2026 | Binary reminder and acceptance edition",
        "Is it clear that simulated policy-output pre-verification and held-out-video effectiveness testing occur after elicitation, not during the Probe?":
            "Is it clear that simulated rule pre-verification and held-out cue-level acceptance testing occur after elicitation, not during the Probe?",
        "Shared scale for Q1-Q6\n1 Use for the current task without interruption | 2 Handle silently and avoid unnecessary detail | 3 Give a brief reminder | 4 Ask before using it | 5 Do not use it unless explicitly requested":
            "Shared binary reminder scale for Q1-Q6\n1 Do not show a privacy reminder | 2 Show a brief privacy reminder",
        "How should reminders work when the VLM can infer sensitive information that is not directly visible?":
            "When a visual cue supports a sensitive inference, should the assistant show a brief privacy reminder?",
        "What balance should the system use between missed risks and unnecessary reminders?":
            "How broad should the trigger be when a rule otherwise calls for a reminder: only the most likely or serious cues, or any plausible cue?",
        "What should happen when the VLM is unsure whether a cue is sensitive?":
            "When the assistant is unsure whether a cue is privacy-sensitive, should it show a brief privacy reminder?",
        "What should happen when sensitive information is visible but unnecessary for the current task?":
            "When privacy-sensitive content is visible but not needed for the current task, should the assistant show a brief privacy reminder?",
        "B. Preferred action for similar content: no intervention / silent handling / brief reminder / ask before use / avoid unless explicitly requested.":
            "B. Reminder decision for similar content: do not show a privacy reminder / show a brief privacy reminder.",
        "awareness status, and preferred action": "awareness status, and binary reminder preference",
        "Do the awareness-status and preferred-action items remain valid and sufficiently brief across every present internal category?":
            "Do the awareness-status and binary reminder-preference items remain valid and sufficiently brief across every present internal category?",
        "Primary outcome: exact preferred-action alignment under counterbalanced policy conditions":
            "Primary outcome: immediate acceptance rating (1–5) after each held-out cue decision",
        "Ordinal action error": "Binary reminder-decision alignment",
        "Over-reminder and under-protection rates": "False-reminder and missed-reminder rates",
        "Use short, manually verified held-out egocentric video windows with temporally emerging cues":
            "Use short, manually verified held-out windows; interpret outcomes as immediate cue-level responses only",
        "Workflow and held-out evaluation:": "Workflow and immediate acceptance evaluation:",
    }
    for paragraph in document.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    table = document.tables[1]
    table.cell(1, 1).text = "Ten direct reminder questions; Q1–Q6 share one binary response scale."
    table.cell(4, 1).text = (
        "Generic versus Q10-only versus full VPrivCal binary policies on manually verified held-out cue windows; "
        "participants rate immediate acceptance and state the preferred reminder decision."
    )
    table.cell(5, 1).text = (
        "Keep/revise/remove decisions for Q10, Probe, Hint, stimuli, binary reminder rules, and cue-level acceptance measures; pilot-readiness vote."
    )

    document.tables[4].cell(0, 0).text = (
        "Shared binary reminder scale for Q1–Q6\n"
        "1 Do not show a privacy reminder | 2 Show a brief privacy reminder"
    )

    evaluation = document.tables[9]
    evaluation.cell(1, 0).text = "Primary outcome: immediate acceptance rating (1–5) after each held-out cue decision"
    evaluation.cell(2, 0).text = "Binary reminder-decision alignment"
    evaluation.cell(3, 0).text = "False-reminder and missed-reminder rates"
    evaluation.cell(4, 0).text = "Separate detector validity from privacy relevance and reminder acceptance"
    evaluation.cell(5, 0).text = (
        "Use short, manually verified held-out windows with emerging cues; do not claim long-term behavior or effectiveness"
    )
    evaluation.cell(6, 0).text = (
        "Simulate binary rule outputs before the pilot; label them non-empirical and keep them separate from participant acceptance evidence"
    )

    for table in document.tables:
        set_table_geometry(table)

    for shape in document.inline_shapes:
        if shape.width > Inches(7):
            ratio = shape.height / shape.width
            shape.width = Inches(7)
            shape.height = int(shape.width * ratio)

    document.core_properties.title = "VPrivCal Expert Review – Binary Reminder and Acceptance Edition"
    document.core_properties.subject = "Expert review of binary privacy-reminder calibration and immediate held-out acceptance"
    document.core_properties.comments = "Version 6.2: printable table geometry and binary reminder outcomes."
    document.save(OUTPUT)

    replacements_by_media = {
        "word/media/image1.png": workflow_path.read_bytes(),
        "word/media/image3.png": comparison_path.read_bytes(),
    }
    temp_output = OUTPUT.with_suffix(".tmp.docx")
    with zipfile.ZipFile(OUTPUT, "r") as source_zip, zipfile.ZipFile(temp_output, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            target_zip.writestr(item, replacements_by_media.get(item.filename, source_zip.read(item.filename)))
    temp_output.replace(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
