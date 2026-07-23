"""Create the v6.4 expert workbook for the five-point cross-cutting policy."""

from __future__ import annotations

import os
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont

from update_expert_workshop_binary_reminders import draw_multiline


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source materials" / "VPrivCal_Expert_Workshop_Interface_Aligned_v6_3.docx"
OUTPUT = ROOT / "source materials" / "VPrivCal_Expert_Workshop_Interface_Aligned_v6_4.docx"
ASSETS = ROOT / ".docx_qa" / "v6_4_assets"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    filename = "arialbd.ttf" if bold else "arial.ttf"
    return ImageFont.truetype(str(Path(os.environ["WINDIR"]) / "Fonts" / filename), size)


def make_workflow(path: Path) -> None:
    image = Image.new("RGB", (2400, 760), "white")
    draw = ImageDraw.Draw(image)
    navy, gray = "#173F5F", "#66758A"
    draw.text((80, 42), "VPrivCal workflow: category triggers + agreement thresholds", font=font(45, True), fill=navy)
    draw.text(
        (80, 104),
        "Expert-verified cues combine with participant thresholds; reminder presentation stays fixed.",
        font=font(24),
        fill=gray,
    )
    boxes = [
        ("1", "VPrivCal-Q10", "Set six category triggers and four five-point agreement thresholds.", "~2 min", "#E8F0F7", navy),
        ("2", "VPrivCal-Probe", "Review three contexts and correct category trigger preferences after seeing evidence.", "~3-5 min", "#F0EAFB", "#6E43B9"),
        ("3", "Verified cue gate", "Two experts confirm category and label presence-only or sensitive-detail exposure.", "Locked before study", "#E6F3EF", "#27836C"),
        ("4", "Binary decision", "The rule produces NO REMINDER or SHOW REMINDER using one fixed presentation.", "Deterministic", "#FFF0E7", "#C65312"),
        ("5", "Acceptance", "Collect cue awareness, 1-5 acceptance, preferred decision, false and missed reminders.", "Primary outcomes", "#E7F3F6", "#18788A"),
    ]
    left, top, width, height, gap = 70, 195, 405, 385, 65
    for index, (number, title, body, footer, bg, accent) in enumerate(boxes):
        x = left + index * (width + gap)
        draw.rounded_rectangle((x, top, x + width, top + height), radius=24, fill=bg, outline=accent, width=5)
        draw.ellipse((x + 24, top + 25, x + 86, top + 87), fill=accent)
        draw.text((x + 55, top + 56), number, font=font(27, True), fill="white", anchor="mm")
        draw.text((x + 105, top + 38), title, font=font(27, True), fill=accent)
        draw_multiline(draw, (x + 26, top + 128), body, width - 52, font(23), "#26384A", 9)
        draw.text((x + 26, top + height - 55), footer, font=font(23, True), fill=accent)
        if index < len(boxes) - 1:
            arrow_x = x + width + 12
            center_y = top + height // 2
            draw.polygon(
                [(arrow_x, center_y - 19), (arrow_x + 36, center_y), (arrow_x, center_y + 19)],
                fill=navy,
            )
            draw.rectangle((arrow_x - 12, center_y - 7, arrow_x + 4, center_y + 7), fill=navy)
    draw.rounded_rectangle((85, 625, 2315, 710), radius=18, fill="#F3F6F8", outline="#C6D1DB", width=3)
    draw.text(
        (1200, 667),
        "The study tests calibration under controlled expert-verified detections, not end-to-end VLM accuracy.",
        font=font(25, True),
        fill=navy,
        anchor="mm",
    )
    image.save(path, quality=96)


def make_comparison(path: Path) -> None:
    image = Image.new("RGB", (1800, 650), "white")
    draw = ImageDraw.Draw(image)
    navy, gray = "#173F5F", "#66758A"
    draw.text((65, 35), "Held-out comparison under verified detections", font=font(43, True), fill=navy)
    draw.text(
        (65, 91),
        "All conditions receive the same locked cue labels and the same reminder presentation.",
        font=font(24),
        fill=gray,
    )
    boxes = [
        ("Generic policy", "One fixed trigger threshold for every participant.", "#E8F0F7", "#173F5F"),
        ("Q10-only policy", "Category priors plus five-point cross-cutting agreement thresholds.", "#F0EAFB", "#6E43B9"),
        ("Full VPrivCal policy", "Q10 priors corrected by category-by-context Probe responses.", "#E6F3EF", "#27836C"),
    ]
    width, height, top, gap = 500, 265, 175, 85
    for index, (title, body, bg, accent) in enumerate(boxes):
        x = 65 + index * (width + gap)
        draw.rounded_rectangle((x, top, x + width, top + height), radius=24, fill=bg, outline=accent, width=5)
        draw.text((x + 28, top + 34), title, font=font(30, True), fill=accent)
        draw_multiline(draw, (x + 28, top + 105), body, width - 56, font(25), "#26384A", 10)
    draw.rounded_rectangle((65, 490, 1735, 603), radius=20, fill="#FFF5EA", outline="#D37A26", width=4)
    draw.text((900, 530), "Each cue still produces SHOW REMINDER or NO REMINDER.", font=font(27, True), fill="#8B4712", anchor="mm")
    draw.text((900, 568), "Compare immediate acceptance and cue awareness without a presentation confound.", font=font(24), fill="#8B4712", anchor="mm")
    image.save(path, quality=96)


def replace_media(docx_path: Path, workflow: Path, comparison: Path) -> None:
    replacements = {
        "word/media/image1.png": workflow.read_bytes(),
        "word/media/image3.png": comparison.read_bytes(),
    }
    temporary = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            target.writestr(item, replacements.get(item.filename, source.read(item.filename)))
    temporary.replace(docx_path)


def set_paragraph(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph content while preserving its paragraph style."""
    paragraph.text = text


def insert_after(paragraph: Paragraph, text: str, style_source: Paragraph) -> Paragraph:
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    inserted = Paragraph(node, paragraph._parent)
    inserted.style = style_source.style
    inserted.text = text
    return inserted


def set_cell(table, row: int, column: int, text: str) -> None:
    table.cell(row, column).text = text


def replace_exact_paragraph(document: Document, old: str, new: str) -> None:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph matching {old!r}; found {len(matches)}")
    set_paragraph(matches[0], new)


def main() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    workflow_path = ASSETS / "workflow_five_point.png"
    comparison_path = ASSETS / "comparison_five_point.png"
    make_workflow(workflow_path)
    make_comparison(comparison_path)
    document = Document(SOURCE)

    set_paragraph(
        document.paragraphs[6],
        "Version 6.4 | July 2026 | Five-point cross-cutting agreement edition",
    )

    # Add the requested brief introduction and background before the workflow.
    anchor = document.paragraphs[9]
    heading_style_source = document.paragraphs[11]
    body_style_source = document.paragraphs[9]
    anchor = insert_after(anchor, "Study introduction and background", heading_style_source)
    anchor = insert_after(
        anchor,
        (
            "VPrivCal is a brief cold-start calibration method for egocentric "
            "vision-language assistants. Six category-specific trigger questions, four "
            "cross-cutting agreement items, and a three-scene Probe produce an auditable "
            "binary policy: show a standardized privacy reminder or do not show it."
        ),
        body_style_source,
    )
    insert_after(
        anchor,
        (
            "Egocentric video can continuously expose identifying, household, bystander, "
            "and contextual information. A fixed policy may over-warn some users and "
            "under-protect others. This review therefore examines whether the five-point "
            "risk thresholds, expert-verified cues, Probe corrections, and held-out "
            "acceptance measures are understandable and suitable for cognitive testing."
        ),
        body_style_source,
    )

    set_cell(
        document.tables[0],
        0,
        0,
        (
            "Decision / research question\nAre the three-option category triggers, "
            "five-point cross-cutting agreement questions, expert-verified exposure labels, "
            "Probe interface, and held-out acceptance evaluation sufficiently clear, "
            "low-burden, and testable?"
        ),
    )

    overview = document.tables[1]
    set_cell(
        overview,
        1,
        1,
        (
            "Ten direct policy questions: Q1-Q6 share one three-option category trigger "
            "scale; Q7-Q10 share one five-point agreement scale with distinct risk thresholds."
        ),
    )
    set_cell(
        overview,
        4,
        1,
        (
            "Generic versus Q10-only versus full VPrivCal under the same expert-verified cue "
            "set and standardized reminder; participants report cue awareness, immediate "
            "acceptance, and preferred binary decision. Q7 is the general show/hide score; "
            "Q10 is retained provisionally so its relationship with acceptance can be tested."
        ),
    )
    set_cell(
        overview,
        5,
        1,
        (
            "Review the category trigger labels, four agreement statements, distinct "
            "likelihood-plus-severity thresholds, exposure annotation, Probe burden, "
            "deterministic rule, controlled comparison, and pilot readiness."
        ),
    )

    scale_callout = document.tables[4]
    set_cell(
        scale_callout,
        0,
        0,
        (
            "Shared three-option category trigger scale for Q1-Q6\n"
            "Do not show reminders for this category | Show reminders only when identifying "
            "or sensitive details are exposed | Show reminders whenever this verified "
            "category is present\n\n"
            "Shared five-point agreement scale for Q7-Q10\n"
            "Strongly disagree | Disagree | Neither agree nor disagree | Agree | Strongly agree\n"
            "Distinct rule thresholds (minimum likelihood + severity score): "
            "11 [never] | 9 | 7 | 5 | 2 [any verified cue]. The observed score ranges from "
            "2 to 10. Q7 applies generally to every verified privacy threat; Q8, Q9, and Q10 "
            "apply when the cue is inferred, uncertain, or task-irrelevant."
        ),
    )
    for paragraph in scale_callout.cell(0, 0).paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        for run in paragraph.runs:
            run.font.size = Pt(8.5)

    # Item-specific prompts: all four use the same agreement question format.
    replacements = {
        "Q7. Inferred risks": "Q7. General reminder sensitivity",
        "When a visual cue supports a sensitive inference, should the assistant show a brief privacy reminder?": (
            "How much do you agree with the following statement?\n"
            '"In general, the assistant should show detected privacy threats to the user."'
        ),
        "Q8. Other verified privacy categories": "Q8. Inferred risks",
        "If experts verify a privacy category not covered above, when should the assistant show a reminder?": (
            "How much do you agree with the following statement?\n"
            '"The assistant should show a privacy reminder when a visual cue supports a '
            'sensitive inference."'
        ),
        "When the assistant is unsure whether a cue is privacy-sensitive, should it show a brief privacy reminder?": (
            "How much do you agree with the following statement?\n"
            '"The assistant should show a privacy reminder when it is uncertain whether a '
            'detected visual cue is privacy-sensitive."'
        ),
        "When privacy-sensitive content is visible but not needed for the current task, should the assistant show a brief privacy reminder?": (
            "How much do you agree with the following statement?\n"
            '"The assistant should show a privacy reminder when privacy-sensitive content is '
            'visible but not needed for the current task."'
        ),
    }
    for old, new in replacements.items():
        replace_exact_paragraph(document, old, new)

    evaluation_grid = document.tables[5]
    set_cell(evaluation_grid, 8, 1, "General reminder sensitivity")

    probe_review = document.tables[8]
    set_cell(
        probe_review,
        5,
        0,
        (
            "Are the awareness-status item, three-option category trigger, and five-point "
            "cross-cutting agreement items distinct, understandable, and sufficiently brief?"
        ),
    )

    held_out_review = document.tables[9]
    set_cell(
        held_out_review,
        6,
        0,
        (
            "Treat simulations as software checks and short-video outcomes as immediate only; "
        "prespecify whether Q7 and Q10 values predict acceptance before considering later "
            "Q10 removal."
        ),
    )

    # Avoid a repeated table header stranded at the foot of the stimuli page.
    page_break = OxmlElement("w:p")
    page_break_properties = OxmlElement("w:pPr")
    page_break_before = OxmlElement("w:pageBreakBefore")
    page_break_properties.append(page_break_before)
    page_break.append(page_break_properties)
    document.tables[7]._tbl.addprevious(page_break)

    # Keep each exposure-codebook row intact instead of splitting one across pages.
    for row in document.tables[10].rows:
        row_properties = row._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))

    document.core_properties.title = "VPrivCal Expert Review Workbook v6.4"
    document.core_properties.subject = (
        "Expert review of five-point cross-cutting agreement thresholds and binary reminders"
    )
    document.core_properties.comments = (
        "Updated from v6.3: unified Q7-Q10 agreement format, Q7 general sensitivity, "
        "distinct five-level risk thresholds, provisional Q10 retention, and study background."
    )
    document.save(OUTPUT)
    replace_media(OUTPUT, workflow_path, comparison_path)
    print(OUTPUT)


if __name__ == "__main__":
    main()
