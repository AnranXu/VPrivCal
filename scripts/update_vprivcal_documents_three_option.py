"""Create three-option VPrivCal research and expert-review documents with safe table geometry."""

from __future__ import annotations

import os
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont

from update_expert_workshop_binary_reminders import draw_multiline, set_table_geometry


ROOT = Path(__file__).resolve().parents[1]
MATERIALS = ROOT / "source materials"
QA = ROOT / ".docx_qa" / "three_option_assets"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    filename = "arialbd.ttf" if bold else "arial.ttf"
    return ImageFont.truetype(str(Path(os.environ["WINDIR"]) / "Fonts" / filename), size)


def make_workflow(path: Path) -> None:
    image = Image.new("RGB", (2400, 760), "white")
    draw = ImageDraw.Draw(image)
    navy, gray = "#173F5F", "#66758A"
    draw.text((80, 42), "VPrivCal workflow: three-option trigger calibration", font=font(48, True), fill=navy)
    draw.text(
        (80, 104),
        "Expert-verified exposure labels combine with participant thresholds; reminder presentation stays fixed.",
        font=font(24),
        fill=gray,
    )
    boxes = [
        ("1", "VPrivCal-Q10", "Choose never, sensitive-detail only, or whenever verified present for each category.", "~2 min", "#E8F0F7", navy),
        ("2", "VPrivCal-Probe", "Review three contexts and correct the category trigger preference after seeing evidence.", "~3-5 min", "#F0EAFB", "#6E43B9"),
        ("3", "Verified cue gate", "Two experts confirm category and label presence-only or sensitive-detail exposure.", "Locked before study", "#E6F3EF", "#27836C"),
        ("4", "Binary decision", "The rule produces NO REMINDER or SHOW REMINDER using one fixed presentation.", "Deterministic", "#FFF0E7", "#C65312"),
        ("5", "Acceptance", "Collect cue awareness, 1-5 acceptance, preferred decision, false and missed reminders.", "Primary outcomes", "#E7F3F6", "#18788A"),
    ]
    left, top, width, height, gap = 70, 195, 405, 385, 65
    for index, (number, title, body, footer, bg, accent) in enumerate(boxes):
        x = left + index * (width + gap)
        draw.rounded_rectangle((x, top, x + width, top + height), radius=24, fill=bg, outline=accent, width=5)
        draw.ellipse((x + 24, top + 25, x + 86, top + 87), fill=accent)
        draw.text((x + 55, top + 56), number, font=font(27, True), fill="white", anchor="mm")
        draw.text((x + 105, top + 38), title, font=font(27, True), fill=accent)
        draw_multiline(draw, (x + 26, top + 128), body, width - 52, font(23), "#26384A", 9)
        draw.text((x + 26, top + height - 55), footer, font=font(23, True), fill=accent)
        if index < len(boxes) - 1:
            arrow_x = x + width + 12
            center_y = top + height // 2
            draw.polygon(
                [(arrow_x, center_y - 19), (arrow_x + 36, center_y), (arrow_x, center_y + 19)],
                fill=navy,
            )
            draw.rectangle((arrow_x - 12, center_y - 7, arrow_x + 4, center_y + 7), fill=navy)
    draw.rounded_rectangle((85, 625, 2315, 710), radius=18, fill="#F3F6F8", outline="#C6D1DB", width=3)
    draw.text(
        (1200, 667),
        "The study tests calibration under controlled expert-verified detections, not end-to-end VLM accuracy.",
        font=font(25, True),
        fill=navy,
        anchor="mm",
    )
    image.save(path, quality=96)


def make_comparison(path: Path) -> None:
    image = Image.new("RGB", (1800, 650), "white")
    draw = ImageDraw.Draw(image)
    navy, gray = "#173F5F", "#66758A"
    draw.text((65, 35), "Held-out comparison under verified detections", font=font(43, True), fill=navy)
    draw.text(
        (65, 91),
        "All conditions receive the same locked cue labels and the same reminder presentation.",
        font=font(24),
        fill=gray,
    )
    boxes = [
        ("Generic policy", "One fixed trigger threshold for every participant.", "#E8F0F7", "#173F5F"),
        ("Q10-only policy", "Three-option category priors plus binary cross-cutting preferences.", "#F0EAFB", "#6E43B9"),
        ("Full VPrivCal policy", "Q10 priors corrected by category-by-context Probe responses.", "#E6F3EF", "#27836C"),
    ]
    width, height, top, gap = 500, 265, 175, 85
    for index, (title, body, bg, accent) in enumerate(boxes):
        x = 65 + index * (width + gap)
        draw.rounded_rectangle((x, top, x + width, top + height), radius=24, fill=bg, outline=accent, width=5)
        draw.text((x + 28, top + 34), title, font=font(30, True), fill=accent)
        draw_multiline(draw, (x + 28, top + 105), body, width - 56, font(25), "#26384A", 10)
    draw.rounded_rectangle((65, 490, 1735, 603), radius=20, fill="#FFF5EA", outline="#D37A26", width=4)
    draw.text(
        (900, 530),
        "Each cue still produces SHOW REMINDER or NO REMINDER.",
        font=font(27, True),
        fill="#8B4712",
        anchor="mm",
    )
    draw.text(
        (900, 568),
        "Compare immediate acceptance and cue awareness without a presentation confound.",
        font=font(24),
        fill="#8B4712",
        anchor="mm",
    )
    image.save(path, quality=96)


def replace_text(document: Document, replacements: dict[str, str]) -> None:
    paragraphs = list(document.paragraphs)
    paragraphs.extend(
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    for paragraph in paragraphs:
        updated = paragraph.text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != paragraph.text:
            paragraph.text = updated


def style_tables(document: Document) -> None:
    for table in document.tables:
        set_table_geometry(table)
    for shape in document.inline_shapes:
        if shape.width > Inches(7):
            ratio = shape.height / shape.width
            shape.width = Inches(7)
            shape.height = int(shape.width * ratio)


EXPOSURE_CODEBOOK_ROWS = [
    (
        "Biometric data",
        "Distant, occluded, or otherwise non-recognizable person/body feature.",
        "A clear face or distinctive body feature that supports recognition or re-identification.",
    ),
    (
        "Children images",
        "A child is present without a recognizable identity or linkable school/location detail.",
        "A recognizable child, readable name, school affiliation, home link, or other identity/location detail.",
    ),
    (
        "PII",
        "A document, screen, badge, plate, or device is present but identifying content is unreadable.",
        "A readable name, address, account/contact number, email, plate, patient ID, or comparable identifier.",
    ),
    (
        "Legal sensitivity",
        "A potentially sensitive object or setting is present without person-specific sensitive content.",
        "Explicit content, a visible unlawful/high-risk act, or person-linked legal/safety information.",
    ),
    (
        "Personal life",
        "A generic home, routine, relationship, or personal object is visible without a linkable private fact.",
        "An intimate activity, precise location, relationship/routine, or household detail is exposed and linkable.",
    ),
    (
        "Background individuals",
        "An incidental person is distant or non-recognizable.",
        "A bystander is recognizable or linked to a readable name, workplace, health context, or other private detail.",
    ),
]


def insert_exposure_codebook(
    document: Document,
    before_heading: str,
    section_heading: str,
    review_prompt: str | None = None,
) -> None:
    target = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == before_heading)
    heading = target.insert_paragraph_before(section_heading)
    heading.style = "Heading 2"
    target.insert_paragraph_before(
        "Experts first confirm category membership. Then ask whether the visible cue identifies or links a person, "
        "or exposes a private fact about a person or household. If yes, label SENSITIVE_DETAIL_EXPOSED; "
        "otherwise label PRESENCE_ONLY. Disagreements are adjudicated or excluded before the cue JSON is locked."
    )

    table = document.add_table(rows=1, cols=3)
    table.style = document.tables[1].style
    headers = ("Category", "PRESENCE_ONLY anchor", "SENSITIVE_DETAIL_EXPOSED anchor")
    for column, value in enumerate(headers):
        table.cell(0, column).text = value
    for values in EXPOSURE_CODEBOOK_ROWS:
        cells = table.add_row().cells
        for column, value in enumerate(values):
            cells[column].text = value
    for column, width in enumerate((1.2, 2.85, 2.95)):
        for cell in table.columns[column].cells:
            cell.width = Inches(width)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5 if row_index else 9)
                    if row_index == 0:
                        run.font.bold = True
    target._p.addprevious(table._tbl)
    if review_prompt:
        target.insert_paragraph_before(review_prompt)


def replace_media(docx_path: Path, workflow: Path, comparison: Path) -> None:
    replacements = {
        "word/media/image1.png": workflow.read_bytes(),
        "word/media/image3.png": comparison.read_bytes(),
    }
    temporary = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            target.writestr(item, replacements.get(item.filename, source.read(item.filename)))
    temporary.replace(docx_path)


def update_interface_workshop(workflow: Path, comparison: Path) -> Path:
    source = MATERIALS / "VPrivCal_Expert_Workshop_Interface_Aligned_v6_2.docx"
    output = MATERIALS / "VPrivCal_Expert_Workshop_Interface_Aligned_v6_3.docx"
    document = Document(source)
    replace_text(document, {
        "Version 6.2 | July 2026 | Binary reminder and acceptance edition":
            "Version 6.3 | July 2026 | Three-option trigger and acceptance edition",
        "revised binary reminder questions": "revised three-option trigger questions",
        "binary reminder preference": "three-option reminder-trigger preference",
        "binary reminder-preference": "three-option reminder-trigger",
        "Q8. Reminder sensitivity": "Q8. Other verified privacy categories",
        "How broad should the trigger be when a rule otherwise calls for a reminder: only the most likely or serious cues, or any plausible cue?":
            "If experts verify a privacy category not covered above, when should the assistant show a reminder?",
        "Workflow and immediate acceptance evaluation:": "Workflow, expert verification, and immediate acceptance evaluation:",
    })
    document.tables[0].cell(0, 0).text = (
        "Decision / research question\nAre the three-option trigger questions, expert-verified exposure labels, "
        "Probe interface, and held-out acceptance evaluation sufficiently clear, low-burden, and testable?"
    )
    method = document.tables[1]
    method.cell(1, 1).text = "Ten direct policy questions; Q1-Q6 share one three-option trigger scale."
    method.cell(4, 1).text = (
        "Generic versus Q10-only versus full VPrivCal under the same expert-verified cue set and standardized reminder; "
        "participants report cue awareness, immediate acceptance, and preferred binary decision."
    )
    method.cell(5, 1).text = (
        "Review the three trigger labels, exposure annotation, Probe burden, deterministic rule, controlled comparison, and pilot readiness."
    )
    document.tables[3].cell(3, 0).text = (
        "Is it clear that experts lock each cue as presence-only or sensitive-detail exposed before the participant policy runs?"
    )
    document.tables[4].cell(0, 0).text = (
        "Shared three-option trigger scale for Q1-Q6\n"
        "1 Do not show reminders for this category | 2 Show reminders only when identifying or sensitive details are exposed | "
        "3 Show reminders whenever this verified category is present"
    )
    document.tables[5].cell(8, 1).text = "Other verified privacy categories"
    document.tables[6].cell(0, 0).text = (
        "Two questions per internal category-scene pair\n"
        "A. Before the highlight: already aware / noticed but did not consider privacy / did not notice or realize VLM capability / not a concern.\n"
        "B. Preferred trigger: never remind / remind only for identifying or sensitive-detail exposure / remind whenever the verified category is present.\n"
        "Category identities remain hidden; participants see neutral linked-content labels."
    )
    document.tables[8].cell(5, 0).text = (
        "Are the awareness-status and three-option trigger items distinct, understandable, and sufficiently brief across every category?"
    )
    evaluation = document.tables[9]
    evaluation.cell(1, 0).text = "Primary outcome: immediate acceptance rating (1-5) after each standardized reminder decision"
    evaluation.cell(2, 0).text = "Cue-awareness difference between generic and VPrivCal policies"
    evaluation.cell(3, 0).text = "Binary decision alignment, false-reminder rate, and missed-reminder rate"
    evaluation.cell(4, 0).text = "Two-expert verification of category and exposure level before study inclusion"
    evaluation.cell(5, 0).text = "Balance presence-only and sensitive-detail-exposed held-out cues where feasible"
    evaluation.cell(6, 0).text = "Treat simulations as software checks and short-video outcomes as immediate only"
    insert_exposure_codebook(
        document,
        "8. Final decision",
        "7.1 Exposure-label codebook review",
        "Reviewer judgment: Can two independent experts apply these anchors consistently?  "
        "□ Strong  □ Revise  □ Reject    Comment: ____________________________________________",
    )
    style_tables(document)
    document.core_properties.title = "VPrivCal Expert Review - Three-Option Trigger Edition"
    document.core_properties.subject = "Expert review of three trigger preferences and verified-exposure rules"
    document.core_properties.comments = "Version 6.3; explicit printable table widths."
    document.save(output)
    replace_media(output, workflow, comparison)
    return output


def update_expert_workshop(workflow: Path, comparison: Path) -> Path:
    source = MATERIALS / "VPrivCal_Expert_Workshop_Updated_v6.docx"
    output = MATERIALS / "VPrivCal_Expert_Workshop_Updated_v7.docx"
    document = Document(source)
    replace_text(document, {
        "preferred future action": "preferred reminder trigger",
        "preferred-action": "preferred-trigger",
        "Q8. Reminder sensitivity": "Q8. Other verified privacy categories",
        "What balance should the system use between missed risks and unnecessary reminders?":
            "If experts verify a privacy category not covered above, when should the assistant show a reminder?",
        "Workflow and held-out evaluation:": "Workflow, expert verification, and held-out evaluation:",
    })
    document.paragraphs[6].text = "Version 7 | July 2026 | Three-option trigger edition"
    document.tables[0].cell(0, 0).text = (
        "Decision / research question\nAre the three-option trigger questions, expert-verified exposure labels, "
        "Probe workflow, and held-out evaluation ready for cognitive testing?"
    )
    method = document.tables[1]
    method.cell(1, 1).text = "Ten direct policy questions; Q1-Q6 share one three-option trigger scale."
    method.cell(4, 1).text = "Generic versus Q10-only versus full VPrivCal on the same locked expert-verified held-out cues."
    method.cell(5, 1).text = "Keep/revise/remove decisions for the questions, exposure rule, Probe, stimuli, and evaluation."
    document.tables[4].cell(0, 0).text = (
        "Shared three-option trigger scale for Q1-Q6\n"
        "1 Do not show reminders for this category | 2 Show reminders only when identifying or sensitive details are exposed | "
        "3 Show reminders whenever this verified category is present"
    )
    document.tables[5].cell(8, 1).text = "Other verified privacy categories"
    document.tables[6].cell(0, 0).text = (
        "Two participant questions per category-image pair\n"
        "A. Awareness status: already aware / noticed without considering privacy / did not notice or realize VLM capability / not a concern.\n"
        "B. Preferred trigger: never remind / sensitive-detail exposure only / whenever the verified category is present."
    )
    review = document.tables[8]
    review.cell(2, 0).text = "Does the three-option preferred-trigger item map deterministically to the reminder policy?"
    review.cell(3, 0).text = "Is it valid and sufficiently brief to ask the trigger question for every category present?"
    review.cell(4, 0).text = "Could expert-verified evidence still create agreement or demand effects?"
    evaluation = document.tables[9]
    evaluation.cell(1, 0).text = "Primary outcome: immediate reminder-decision acceptance"
    evaluation.cell(2, 0).text = "Privacy-cue awareness under generic versus personalized triggering"
    evaluation.cell(3, 0).text = "Binary alignment, false reminders, and missed reminders"
    evaluation.cell(4, 0).text = "Two-expert verification of detection category and exposure level"
    evaluation.cell(5, 0).text = "Use the same standardized reminder presentation in all conditions"
    evaluation.cell(6, 0).text = "Counterbalanced policy evaluation on locked held-out cues"
    insert_exposure_codebook(
        document,
        "8. Final decision",
        "7.1 Exposure-label codebook review",
        "Reviewer judgment: Can two independent experts apply these anchors consistently?  "
        "□ Strong  □ Revise  □ Reject    Comment: ____________________________________________",
    )
    style_tables(document)
    document.core_properties.title = "VPrivCal Expert Workshop - Three-Option Trigger Edition"
    document.core_properties.comments = "Version 7; explicit printable table widths."
    document.save(output)
    replace_media(output, workflow, comparison)
    return output


def remove_table_rows(table, keep_rows: int) -> None:
    while len(table.rows) > keep_rows:
        row = table.rows[-1]
        table._tbl.remove(row._tr)


def update_research_plan(workflow: Path, comparison: Path) -> Path:
    source = MATERIALS / "VPrivCal_Full_Research_Plan_Updated_v6.docx"
    output = MATERIALS / "VPrivCal_Full_Research_Plan_Updated_v7.docx"
    document = Document(source)
    replace_text(document, {
        "awareness and preferred-action questions": "awareness and preferred-trigger questions",
        "preferred-action": "preferred-trigger",
        "preferred future action": "preferred reminder trigger",
        "suppression, handle silently, remind, ask, or avoid": "a no-reminder or standardized-reminder decision",
        "Q8. Reminder sensitivity": "Q8. Other verified privacy categories",
        "Ordinal error and preferred action": "Binary decision alignment and preferred decision",
        "salience, direct versus inferred cue, and temporal dependence":
            "expert-confirmed exposure level, inference status, and temporal dependence",
    })
    document.paragraphs[1].text = (
        "A Brief Three-Option Preference-Elicitation and Awareness-Confirmation Method for "
        "User-Aligned Privacy Reminders in Vision-Language Assistants"
    )
    document.paragraphs[6].text = (
        "Version 7 | July 2026 | Three-option trigger and controlled-evaluation edition"
    )
    document.paragraphs[9].text = (
        "VPrivCal is a brief cold-start calibration method for egocentric VLM assistants. It asks ten direct policy questions, "
        "then uses three private, public, and semi-public scenes. Q1-Q6 and Probe use three trigger preferences: never remind, "
        "remind only when identifying or sensitive details are exposed, or remind whenever the expert-verified category is present."
    )
    document.paragraphs[10].text = (
        "The calibration is not the effectiveness outcome. Under a controlled locked cue set, VPrivCal is supported only if it "
        "improves immediate reminder-decision acceptance and privacy-cue awareness relative to the generic policy without increasing missed reminders."
    )
    document.paragraphs[39].text = (
        "The fixed VLM supplies candidates, but two independent experts verify every included cue and assign PRESENCE_ONLY or "
        "SENSITIVE_DETAIL_EXPOSED before the study. A transparent policy layer then produces no reminder or one standardized reminder. "
        "The study does not test end-to-end detector accuracy."
    )
    document.paragraphs[42].text = (
        "Recommended conservative guardrail: apply any approved safety floor only after the participant trigger is resolved, record every override, "
        "and review whether floors erase meaningful three-option personalization."
    )
    document.paragraphs[46].text = (
        "Static calibration and dynamic evaluation remain intentionally different. Held-out clips use locked expert-verified detections and the same "
        "reminder presentation in every condition. Include both exposure labels when feasible so the middle trigger can be evaluated."
    )
    document.paragraphs[61].text = (
        "Use a 7-point agreement scale for overall policy experience, while retaining the cue-level 1-5 acceptance item. Analyze acceptance, awareness, "
        "burden, protection/control, and adoption separately unless psychometric evidence supports a composite."
    )
    document.paragraphs[64].text = (
        "Binary decision alignment: mixed-effects logistic model or participant-level paired comparison, with exposure level included as a prespecified factor."
    )

    document.tables[0].cell(0, 0).text = (
        "Decision / research question\nCan a brief three-option calibration improve immediate acceptance and privacy-cue awareness "
        "for standardized VLM reminders under controlled expert-verified detections?"
    )
    overview = document.tables[1]
    overview.cell(3, 1).text = (
        "Three egocentric scenes; participants point first, then review every verified category and choose one of three reminder-trigger thresholds."
    )
    overview.cell(4, 1).text = "A transparent user-specific trigger policy with binary reminder output."
    overview.cell(5, 1).text = "Generic versus Q10-only versus full VPrivCal on identical expert-verified held-out cues."
    overview.cell(6, 1).text = "Immediate acceptance, privacy-cue awareness, binary alignment, false reminders, missed reminders, and burden."

    components = document.tables[3]
    components.cell(1, 2).text = "Six category defaults using the three-option trigger scale plus four cross-cutting settings."
    components.cell(2, 2).text = "Awareness status and preferred trigger for every category-image pair."
    components.cell(3, 2).text = "Weighted category priors and context corrections combined with locked exposure labels."
    components.cell(4, 2).text = "Acceptance and awareness for binary reminder decisions on verified held-out cues."

    questions = document.tables[4]
    question_rows = [
        ("RQ1", "Does full VPrivCal increase immediate reminder-decision acceptance over a generic policy?"),
        ("RQ2", "Does full VPrivCal improve privacy-cue awareness without increasing missed reminders?"),
        ("RQ3", "Does the sensitive-detail-only option produce distinct decisions for the two expert-confirmed exposure levels?"),
        ("RQ4", "Does Probe improve trigger alignment beyond Q10 alone across categories and contexts?"),
        ("H1", "Full VPrivCal will produce higher cue-level acceptance than the generic policy."),
        ("H2", "Full VPrivCal will preserve or improve cue awareness while reducing unwanted reminders."),
        ("H3", "Trigger alignment will be higher when held-out exposure level matches the participant's selected threshold."),
    ]
    for index, (identifier, text) in enumerate(question_rows, start=1):
        questions.cell(index, 0).text = identifier
        questions.cell(index, 1).text = text

    document.tables[5].cell(0, 0).text = (
        "Shared question for Q1-Q6\nWhen this type of information appears, when should the assistant show a privacy reminder?\n\n"
        "1 Do not show reminders for this category | 2 Show reminders only when identifying or sensitive details are exposed | "
        "3 Show reminders whenever this verified category is present"
    )

    cross = document.tables[7]
    cross.cell(1, 2).text = "When an expert-verified cue supports a sensitive inference, should the assistant show a reminder?"
    cross.cell(2, 1).text = "Other verified privacy categories"
    cross.cell(2, 2).text = "If experts verify a privacy category not covered above, when should the assistant show a reminder?"
    cross.cell(2, 3).text = "unlisted_category_trigger"
    cross.cell(3, 2).text = "When the VLM is uncertain about an expert-confirmed privacy cue, should the assistant show a reminder?"
    cross.cell(4, 2).text = "When verified privacy-sensitive content is not needed for the current task, should the assistant show a reminder?"

    phases = document.tables[10]
    phases.cell(2, 3).text = "Awareness status and preferred three-option trigger for each category-image pair."
    phases.cell(3, 2).text = "Combine Q10 and Probe trigger levels; resolve against the expert-confirmed exposure label."
    phases.cell(3, 3).text = "Trigger level and resulting binary reminder decision."

    document.tables[12].cell(0, 0).text = (
        "Question B: preferred reminder trigger\nWhen similar content appears in the future, when should the assistant remind you?\n"
        "1 Do not show reminders for this category.\n"
        "2 Show reminders only when identifying or sensitive details are exposed.\n"
        "3 Show reminders whenever this verified category is present."
    )
    document.tables[14].cell(0, 0).text = (
        "Critical interpretation rule\nAwareness status explains what the participant noticed. The three-option response sets a trigger threshold. "
        "Experts, not participants or an unrestricted VLM judgment, assign the held-out exposure label before the study."
    )

    trigger_table = document.tables[15]
    remove_table_rows(trigger_table, 4)
    trigger_table.cell(0, 0).text = "Trigger level"
    trigger_table.cell(0, 1).text = "Participant label"
    trigger_table.cell(0, 2).text = "Deterministic runtime behavior"
    trigger_rows = [
        ("0", "Never remind", "NO_REMINDER for this category."),
        ("1", "Sensitive-detail exposure only", "SHOW_REMINDER only for SENSITIVE_DETAIL_EXPOSED; otherwise NO_REMINDER."),
        ("2", "Whenever verified present", "SHOW_REMINDER for either expert-confirmed exposure level."),
    ]
    for index, values in enumerate(trigger_rows, start=1):
        for column, value in enumerate(values):
            trigger_table.cell(index, column).text = value

    design = document.tables[16]
    design.cell(2, 1).text = "Within-subject comparison of generic, Q10-only, and full VPrivCal; clip-policy assignment counterbalanced."
    design.cell(3, 1).text = "10-12 short clips with locked expert-verified cues; include both exposure levels where feasible."
    design.cell(5, 1).text = "Balance categories and PRESENCE_ONLY versus SENSITIVE_DETAIL_EXPOSED labels; avoid cells too small to interpret."
    design.cell(6, 1).text = "Exclude unresolved or disputed detections from the main effectiveness comparison."
    design.cell(7, 1).text = "Use one fixed VLM; two experts independently verify category and exposure level before locking the cue JSON."

    cue_questions = document.tables[17]
    cue_rows = [
        ("Expert verification gate", "Do both experts confirm the cue and category?", "Agree / adjudicated / excluded"),
        ("Exposure label", "Does the cue merely show category presence or expose identifying/sensitive details?", "PRESENCE_ONLY / SENSITIVE_DETAIL_EXPOSED"),
        ("Cue awareness", "Before the reminder, had you noticed the privacy-relevant content?", "Yes / partly / no"),
        ("Preferred decision", "Should the assistant have shown a reminder for this cue?", "No reminder / show reminder"),
        ("Decision acceptance", "How acceptable was the assistant's reminder decision?", "1 Not at all - 5 Completely"),
    ]
    for index, values in enumerate(cue_rows, start=1):
        for column, value in enumerate(values):
            cue_questions.cell(index, column).text = value

    document.tables[18].cell(0, 0).text = (
        "Why detector verification is separate\nThe controlled study estimates the effect of VPrivCal under a locked verified cue set. "
        "The VLM candidate error rate and expert agreement must be reported separately from participant acceptance and awareness."
    )

    metrics = document.tables[19]
    metric_rows = [
        ("Immediate acceptance", "Mean cue-level acceptance and proportion rated 4-5.", "Primary"),
        ("Privacy-cue awareness", "Unaided or pre-reminder awareness of the verified cue.", "Primary"),
        ("Binary decision alignment", "System reminder decision equals the participant's preferred decision.", "Primary"),
        ("False-reminder rate", "System reminds when the participant prefers no reminder.", "Burden guardrail"),
        ("Missed-reminder rate", "System does not remind when the participant prefers a reminder.", "Privacy guardrail"),
        ("Exposure-stratified effect", "Outcomes reported separately for both expert-confirmed exposure levels.", "Construct check"),
        ("Probe awareness profile", "Already aware / interpretive gap / perceptual-capability gap / rejected cue.", "Descriptive"),
        ("Reminder frequency", "Number and proportion of verified cues that trigger the standard reminder.", "Dose measure"),
        ("Calibration burden", "Completion time, skipped items, changes, and perceived effort.", "Feasibility"),
        ("User-friendliness", "Relevance, number/timing, interruption, control, and willingness to use.", "Secondary"),
    ]
    for index, values in enumerate(metric_rows, start=1):
        for column, value in enumerate(values):
            metrics.cell(index, column).text = value

    phases_table = document.tables[21]
    phases_table.cell(1, 2).text = "Review three-option wording, exposure codebook, cue agreement, safety floors, and study feasibility."

    insert_exposure_codebook(
        document,
        "6. Personalized policy construction",
        "5.4 Locked exposure-label codebook",
        "Report independent category agreement, exposure-label agreement, adjudications, and exclusions. "
        "If the held-out set does not contain both exposure levels, trigger levels 1 and 2 cannot be distinguished empirically.",
    )
    style_tables(document)
    document.core_properties.title = "VPrivCal Full Research Plan - Three-Option Trigger Edition"
    document.core_properties.subject = "Controlled evaluation with expert-verified detections and exposure labels"
    document.core_properties.comments = "Version 7; explicit printable table widths."
    document.save(output)
    replace_media(output, workflow, comparison)
    return output


def main() -> None:
    QA.mkdir(parents=True, exist_ok=True)
    workflow = QA / "workflow_three_option.png"
    comparison = QA / "comparison_three_option.png"
    make_workflow(workflow)
    make_comparison(comparison)
    outputs = [
        update_interface_workshop(workflow, comparison),
        update_expert_workshop(workflow, comparison),
        update_research_plan(workflow, comparison),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
